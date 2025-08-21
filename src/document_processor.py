"""
Document processing module for the Enterprise Copilot demo.
Handles document ingestion, text extraction, and chunking as per the technical architecture.
"""

import streamlit as st
import PyPDF2
import docx
import io
import re
from typing import List, Dict, Any, Optional
# Try to import ML libraries, fallback to basic implementations
try:
    import tiktoken
    TIKTOKEN_AVAILABLE = True
except ImportError:
    TIKTOKEN_AVAILABLE = False

try:
    from sentence_transformers import SentenceTransformer
    SENTENCE_TRANSFORMERS_AVAILABLE = True
except ImportError:
    SENTENCE_TRANSFORMERS_AVAILABLE = False


class DocumentProcessor:
    """
    Handles document processing including text extraction and chunking.
    Implements sliding window chunking as described in the technical architecture.
    """
    
    def __init__(self, chunk_size: int = 512, chunk_overlap: int = 50):
        """
        Initialize the document processor.
        
        Args:
            chunk_size: Maximum tokens per chunk
            chunk_overlap: Number of overlapping tokens between chunks
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        
        # Initialize tokenizer for token counting
        if TIKTOKEN_AVAILABLE:
            try:
                self.tokenizer = tiktoken.get_encoding("cl100k_base")  # GPT-4 tokenizer
            except Exception:
                self.tokenizer = None
        else:
            # Fallback to approximate token counting
            self.tokenizer = None
            st.info("📏 Using approximate token counting (install tiktoken for precise counts)")
    
    def process_document(self, uploaded_file) -> List[Dict[str, Any]]:
        """
        Process an uploaded document and return chunks.
        
        Args:
            uploaded_file: Streamlit uploaded file object
            
        Returns:
            List of document chunks with metadata
        """
        try:
            # Extract text based on file type
            text = self._extract_text(uploaded_file)
            
            if not text or len(text.strip()) < 10:
                st.error(f"No meaningful text extracted from {uploaded_file.name}")
                return []
            
            # Clean and preprocess text
            cleaned_text = self._clean_text(text)
            
            # Create chunks
            chunks = self._create_chunks(cleaned_text, uploaded_file.name)
            
            return chunks
            
        except Exception as e:
            st.error(f"Error processing {uploaded_file.name}: {str(e)}")
            return []
    
    def _extract_text(self, uploaded_file) -> str:
        """
        Extract text from uploaded file based on file type.
        
        Args:
            uploaded_file: Streamlit uploaded file object
            
        Returns:
            Extracted text content
        """
        file_extension = uploaded_file.name.lower().split('.')[-1]
        
        # Reset file pointer
        uploaded_file.seek(0)
        
        if file_extension == 'pdf':
            return self._extract_pdf_text(uploaded_file)
        elif file_extension == 'docx':
            return self._extract_docx_text(uploaded_file)
        elif file_extension == 'txt':
            return self._extract_txt_text(uploaded_file)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
    
    def _extract_pdf_text(self, uploaded_file) -> str:
        """Extract text from PDF file."""
        try:
            pdf_reader = PyPDF2.PdfReader(uploaded_file)
            text = ""
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                except Exception as e:
                    st.warning(f"Could not extract text from page {page_num + 1}: {str(e)}")
                    continue
            
            return text
            
        except Exception as e:
            raise Exception(f"Error reading PDF: {str(e)}")
    
    def _extract_docx_text(self, uploaded_file) -> str:
        """Extract text from DOCX file."""
        try:
            doc = docx.Document(uploaded_file)
            text = ""
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
            
            # Extract text from tables if any
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text += " | ".join(row_text) + "\n"
            
            return text
            
        except Exception as e:
            raise Exception(f"Error reading DOCX: {str(e)}")
    
    def _extract_txt_text(self, uploaded_file) -> str:
        """Extract text from TXT file."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    uploaded_file.seek(0)
                    content = uploaded_file.read()
                    if isinstance(content, bytes):
                        text = content.decode(encoding)
                    else:
                        text = content
                    return text
                except UnicodeDecodeError:
                    continue
            
            raise Exception("Could not decode file with any supported encoding")
            
        except Exception as e:
            raise Exception(f"Error reading TXT file: {str(e)}")
    
    def _clean_text(self, text: str) -> str:
        """
        Clean and preprocess extracted text.
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned text
        """
        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n', '\n\n', text)
        text = re.sub(r'\s+', ' ', text)
        
        # Remove page markers and similar artifacts
        text = re.sub(r'--- Page \d+ ---', '', text)
        text = re.sub(r'pg\. \d+', '', text)
        
        # Remove excessive punctuation
        text = re.sub(r'\.{3,}', '...', text)
        text = re.sub(r'-{3,}', '---', text)
        
        # Clean up special characters
        text = text.replace('\x00', '')  # Remove null bytes
        text = text.replace('\ufeff', '')  # Remove BOM
        
        return text.strip()
    
    def _create_chunks(self, text: str, document_name: str) -> List[Dict[str, Any]]:
        """
        Create overlapping chunks from the text using sliding window approach.
        
        Args:
            text: Cleaned text to chunk
            document_name: Name of source document
            
        Returns:
            List of chunks with metadata
        """
        chunks = []
        
        # Split text into sentences for better chunk boundaries
        sentences = self._split_into_sentences(text)
        
        if not sentences:
            return []
        
        current_chunk = ""
        current_tokens = 0
        chunk_id = 0
        
        i = 0
        while i < len(sentences):
            sentence = sentences[i]
            sentence_tokens = self._count_tokens(sentence)
            
            # If adding this sentence would exceed chunk size, finalize current chunk
            if current_tokens + sentence_tokens > self.chunk_size and current_chunk:
                # Create chunk
                chunk = {
                    'content': current_chunk.strip(),
                    'document': document_name,
                    'chunk_id': chunk_id,
                    'start_sentence': i - len(current_chunk.split('. ')),
                    'tokens': current_tokens
                }
                chunks.append(chunk)
                chunk_id += 1
                
                # Start new chunk with overlap
                overlap_content = self._get_overlap_content(current_chunk, self.chunk_overlap)
                current_chunk = overlap_content + " " + sentence if overlap_content else sentence
                current_tokens = self._count_tokens(current_chunk)
            else:
                # Add sentence to current chunk
                if current_chunk:
                    current_chunk += " " + sentence
                else:
                    current_chunk = sentence
                current_tokens += sentence_tokens
            
            i += 1
        
        # Add final chunk if it has content
        if current_chunk.strip():
            chunk = {
                'content': current_chunk.strip(),
                'document': document_name,
                'chunk_id': chunk_id,
                'start_sentence': max(0, len(sentences) - len(current_chunk.split('. '))),
                'tokens': current_tokens
            }
            chunks.append(chunk)
        
        return chunks
    
    def _split_into_sentences(self, text: str) -> List[str]:
        """
        Split text into sentences for better chunk boundaries.
        
        Args:
            text: Text to split
            
        Returns:
            List of sentences
        """
        # Simple sentence splitting (could be improved with NLTK or spaCy)
        sentences = re.split(r'(?<=[.!?])\s+', text)
        
        # Filter out very short sentences and clean up
        sentences = [s.strip() for s in sentences if s.strip() and len(s.strip()) > 10]
        
        return sentences
    
    def _count_tokens(self, text: str) -> int:
        """
        Count tokens in text.
        
        Args:
            text: Text to count tokens for
            
        Returns:
            Number of tokens
        """
        if self.tokenizer and TIKTOKEN_AVAILABLE:
            return len(self.tokenizer.encode(text))
        else:
            # Rough approximation: 1 token ≈ 4 characters (or ~0.75 words)
            word_count = len(text.split())
            return max(word_count, len(text) // 4)
    
    def _get_overlap_content(self, text: str, overlap_tokens: int) -> str:
        """
        Get the last part of text for overlap with next chunk.
        
        Args:
            text: Source text
            overlap_tokens: Number of tokens to overlap
            
        Returns:
            Overlap content
        """
        words = text.split()
        
        if len(words) <= overlap_tokens // 4:  # Rough token-to-word ratio
            return text
        
        # Take last N words approximately equal to overlap tokens
        overlap_words = words[-(overlap_tokens // 4):]
        return " ".join(overlap_words)
